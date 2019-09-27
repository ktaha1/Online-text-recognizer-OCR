#!C:\Program Files\Python37\python.exe
# coding: utf-8

import os
import sys
sys.path.append('C:/Users/Taha/AppData/Roaming/Python/Python37/site-packages/')
import cv2 as cv
import pytesseract
from docx import Document
from fpdf import FPDF
import codecs
sys.path.append('/core')
from rotate import dorotate


dorotate('container/input.png')

img = cv.imread('container/outputRotate.png',0)

th2 = cv.adaptiveThreshold(img,255,cv.ADAPTIVE_THRESH_GAUSSIAN_C,cv.THRESH_BINARY,11,4)

out=cv.fastNlMeansDenoising(th2, None, 55, 7, 21);

cv.imwrite('container/outputFinal.png',out)


txt=pytesseract.image_to_string(out,lang=sys.argv[1])


#Text File output
text_file = codecs.open("container/output.txt", "w", "utf-8")
text_file.write(txt)
text_file.close()

# Document file output
document = Document()
document.add_heading('Text extracted:', 0)
document.add_paragraph(txt)
document.save('container/output_doc.docx')


# PDF file output
pdf = FPDF()
pdf.add_page()
pdf.add_font('DejaVu', '', 'core/DejaVuSansCondensed.ttf', uni=True)
pdf.set_font("DejaVu", size=14)
pdf.write(5, txt)
#pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
#pdf.set_font('DejaVu', '', 14)


#bb=txt.encode('utf-8')
#res=bb.decode("utf-8", "replace")

#pdf.multi_cell(350, 15, txt=txt, border = 0)
pdf.output("container/output_pdf.pdf",'F')






